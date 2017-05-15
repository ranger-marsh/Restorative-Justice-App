'''
Takes in a list of values from the database and creates a facesheet.
'''

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def assemble_address(street, apartment, city, state, zip_code):
    address = street.title()
    address += f' APT: {apartment.title()}'
    address += f' {city.title()}, '
    address += state.upper()
    address += ' ' + zip_code
    return address


def parse_row(row_list):
    info = {'case_number': row_list[1],
            'occurred_date': row_list[2],
            'incident_type': row_list[3].title(),
            'age': row_list[5],
            'name':  row_list[7].title(),
            'address': assemble_address(row_list[8], row_list[9],
                                        row_list[10], row_list[11],
                                        row_list[12],
                                        ),
            'DOB':  row_list[13],
            'phone': row_list[14],
            'race': row_list[15].title(),
            'sex': row_list[16].title(),
            'district': row_list[17].title()}
    return info


def district_line(document, district):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'District: {district}').bold = True


def approval_line(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p.add_run('Selection: ').bold = True
    p.add_run('Pass').bold = True
    p.add_run(' | ').bold = True
    p.add_run('Fail').bold = True

    p.add_run().add_break()
    p.add_run('Background: ').bold = True
    p.add_run('Pass').bold = True
    p.add_run(' | ').bold = True
    p.add_run('Fail').bold = True
    p.add_run().add_break()


def case_number_line(document, case_number):
    p = document.add_paragraph()
    p.add_run(f'Case Number: {case_number}')


def name_line(document, name):
    p = document.add_paragraph()
    p.add_run(f'Name: {name}')


def bio_line(document, sex, race, dob, age):
    lines = ['Sex:\t', 'Race:\t', 'DOB:\t', 'Age:\t']
    bio_list = [sex, race, dob, age]
    p = document.add_paragraph()
    for line, bio in zip(lines, bio_list):
        p.add_run(f'{line}{bio}')
        p.add_run().add_break()


def charge_line(document):
    lines = ['Charge Type: State | Municipal',
             'Description:', 'Court Date:', 'Citation#:']
    p = document.add_paragraph()
    for line in lines:
        p.add_run(line)
        p.add_run().add_break()


def phone_line(document, phone):
    p = document.add_paragraph()
    p.add_run(f'Phone: {phone}')
    p.add_run().add_break()
    p.add_run('Email:')


def background_line(document):
    lines = ['Court Records:', 'Out of State Records:',
             'Local Records:', 'Notes:']
    for line in lines:
        p = document.add_paragraph()
        p.add_run(line).bold = True


def last_name_first(name):
    suffix = ['II', 'IV', 'JR', 'SR']
    name_list = name.split()
    name_list.insert(0, name_list.pop())
    if name_list[0][:2].upper() in suffix:
        name_list.insert(0, name_list.pop())
    name = "_".join(name_list)
    return name


def save_facesheet(document, name, dir, district, district_folders):
    name = last_name_first(name)
    if district_folders:
        path = f'{dir}/results/{district}/{name}/{name}.docx'
        if not os.path.isdir(f'{dir}/results/{district}/{name}'):
            os.makedirs(f'{dir}/results/{district}/{name}')
    else:
        path = f'{dir}/results/{name}/{name}.docx'
        if not os.path.isdir(f'{dir}/results/{name}'):
            os.makedirs(f'{dir}/results/{name}')
    document.save(path)


def main():
    pass


if __name__ == '__main__':
    main()
