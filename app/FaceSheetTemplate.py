''' 
Create a face-sheet with information from the sorted Excel file. 
This class assumes that all string formating is done before strings are 
passed to it. The methods are in the order they run and create the document 
from top down. This class is tested functionally by looking at the document it 
creates. 
'''

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class FaceSheetTemplate:

    def __init__(self, district, case_number, name, sex, race, dob, age, address, phone):
        self.document = Document()
        self.district_line(district)
        self.approval_line()
        self.case_number_line(case_number)
        self.name_line(name)
        self.bio_line(sex, race, dob, age)
        self.address_line(address)
        self.phone_line(phone)
        self.charge_line()
        self.background_lines()

    def district_line(self, district):
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run('{}{}'.format('District: ', district)).bold = True

    def approval_line(self):
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p.add_run('Selection: ').bold = True
        p.add_run('Pass').bold = True
        p.add_run(' | ').bold = True
        p.add_run('Fail').bold = True

        p.add_run().add_break()
        p.add_run('Background: ').bold = True
        p.add_run('Pass').bold = True
        p.add_run(' | ').bold = True
        p.add_run('Fail').bold = True
        p.add_run().add_break()

    def case_number_line(self, case_number):
        p = self.document.add_paragraph()
        p.add_run('{}{}'.format('Case Number: ', case_number))

    def name_line(self, name):
        p = self.document.add_paragraph()
        p.add_run('{}{}'.format('Name: ', name))

    def bio_line(self, sex, race, dob, age):
        lines = ['Sex:\t', 'Race:\t', 'DOB:\t', 'Age:\t']
        bio_list = [sex, race, dob, age]
        p = self.document.add_paragraph()
        for line, bio in zip(lines, bio_list):
            p.add_run('{}{}'.format(line, bio))
            p.add_run().add_break()

    def charge_line(self):
        lines = ['Charge Type: State | Municipal',
                 'Description:', 'Court Date:']

        p = self.document.add_paragraph()
        for line in lines:
            p.add_run(line)
            p.add_run().add_break()

    def address_line(self, address):
        p = self.document.add_paragraph()
        p.add_run('{}{}'.format('Address: ', address))

    def phone_line(self, phone):
        p = self.document.add_paragraph()
        p.add_run('{}{}'.format('Phone: ', phone))
        p.add_run().add_break()
        p.add_run('Email:')

    def background_lines(self):
        lines = ['Court Records:', 'Out of State Records:', 'Local Records:', 'Notes:']
        for line in lines:
            p = self.document.add_paragraph()
            p.add_run(line).bold = True

    def save_facesheet(self, path):
        self.document.save(path)

